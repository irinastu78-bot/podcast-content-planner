"""
export_service.py
Экспорт результатов в Word (.docx) через python-docx.
"""

import io
from docx import Document


def summary_to_docx(summary: dict) -> bytes:
    """Конспект статьи -> .docx (байты)."""
    doc = Document()
    doc.add_heading("Article Summary", level=0)
    doc.add_heading("Конспект статьи", level=1)

    if summary.get("title"):
        doc.add_heading(summary["title"], level=2)

    doc.add_heading("Краткое содержание", level=3)
    doc.add_paragraph(summary.get("summary", ""))

    if summary.get("key_points"):
        doc.add_heading("Основные мысли", level=3)
        for kp in summary["key_points"]:
            doc.add_paragraph(kp, style="List Bullet")

    if summary.get("conclusions"):
        doc.add_heading("Выводы", level=3)
        for c in summary["conclusions"]:
            doc.add_paragraph(c, style="List Bullet")

    if summary.get("note"):
        doc.add_heading("Примечание", level=3)
        doc.add_paragraph(summary["note"])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def spoken_text_to_docx(spoken_text, episode_title="", params=None):
    """
    Экспорт текста для озвучки в Word.
    spoken_text — список словарей {speaker, text}; speaker "__block__" — заголовок блока.
    episode_title — название выпуска.
    params — словарь параметров подкаста (для шапки). Может быть None.
    Возвращает bytes готового .docx.
    """
    from io import BytesIO
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    # --- Шапка с технической информацией ---
    p = params or {}
    podcast_name = p.get("podcast_name") or ""
    main_title = podcast_name or episode_title or "Текст для озвучки"
    doc.add_heading(main_title, level=0)

    if episode_title and episode_title != main_title:
        doc.add_heading(f"Выпуск: {episode_title}", level=1)

    info_pairs = [
        ("Длительность", p.get("duration")),
        ("Аудитория", p.get("audience")),
        ("Формат", p.get("format")),
        ("Тон и стиль", p.get("tone")),
        ("Ведущий", p.get("host_name")),
        ("Гость", p.get("guest_name")),
    ]
    info_pairs = [(k, v) for k, v in info_pairs if v]
    if info_pairs:
        for k, v in info_pairs:
            line = doc.add_paragraph()
            run_k = line.add_run(f"{k}: ")
            run_k.bold = True
            line.add_run(str(v))
        doc.add_paragraph("")  # отступ перед текстом

    doc.add_paragraph("—" * 30)

    # --- Тело: реплики и заголовки блоков ---
    for line in spoken_text:
        speaker = (line.get("speaker") or "").strip()
        text = line.get("text", "")
        if speaker == "__block__":
            doc.add_heading(text, level=2)
            continue
        para = doc.add_paragraph()
        if speaker:
            run_s = para.add_run(f"{speaker}: ")
            run_s.bold = True
            run_s.font.size = Pt(12)
            para.add_run(text)
        else:
            para.add_run(text)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def tts_voices_to_docx(voices: dict) -> bytes:
    """Рекомендации по голосам для TTS -> .docx (байты)."""
    doc = Document()
    doc.add_heading("Рекомендации по голосам для озвучки", level=0)
    for v in voices.get("voices", []):
        doc.add_heading(v.get("role", "Голос"), level=2)
        doc.add_paragraph(f"Пол: {v.get('gender', '—')}")
        doc.add_paragraph(f"Возраст: {v.get('age', '—')}")
        doc.add_paragraph(f"Тембр: {v.get('timbre', '—')}")
        doc.add_paragraph(f"Темп: {v.get('pace', '—')}")
        doc.add_paragraph(f"Эмоция: {v.get('emotion', '—')}")
        if v.get("notes"):
            doc.add_paragraph(f"Примечания: {v['notes']}")
    if voices.get("general"):
        doc.add_heading("Общие рекомендации", level=2)
        doc.add_paragraph(voices["general"])
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
